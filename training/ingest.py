"""
Training System — Trade-CLI Phase 2

Knowledge ingestion from PDFs, Markdown, TXT.
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
import logging

logger = logging.getLogger(__name__)


def ingest_document(file_path: str) -> Optional[str]:
    """
    Read document file (PDF, Markdown, TXT).
    
    Args:
        file_path: Path to document
        
    Returns:
        Document text, or None if failed
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        logger.error(f"File not found: {file_path}")
        return None
    
    file_ext = file_path.suffix.lower()
    
    try:
        if file_ext == '.txt':
            return ingest_text(str(file_path))
        elif file_ext == '.md':
            return ingest_markdown(str(file_path))
        elif file_ext == '.pdf':
            return ingest_pdf(str(file_path))
        elif file_ext == '.docx':
            return ingest_docx(str(file_path))
        else:
            logger.error(f"Unsupported file type: {file_ext}")
            return None
    except Exception as e:
        logger.error(f"Ingestion failed for {file_path}: {e}")
        return None


def ingest_text(file_path: str) -> Optional[str]:
    """Read plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info(f"Ingested text file: {file_path} ({len(content)} chars)")
        return content
    except Exception as e:
        logger.error(f"Failed to ingest text: {e}")
        return None


def ingest_markdown(file_path: str) -> Optional[str]:
    """Read markdown file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info(f"Ingested markdown file: {file_path} ({len(content)} chars)")
        return content
    except Exception as e:
        logger.error(f"Failed to ingest markdown: {e}")
        return None


def ingest_pdf(file_path: str) -> Optional[str]:
    """
    Read PDF file.
    Requires: pip install pypdf
    """
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text = ""
        
        for page_num, page in enumerate(reader.pages):
            text += page.extract_text()
            logger.debug(f"Extracted page {page_num + 1}/{len(reader.pages)}")
        
        logger.info(f"Ingested PDF: {file_path} ({len(text)} chars from {len(reader.pages)} pages)")
        return text
    except ImportError:
        logger.error("pypdf not installed. Install with: pip install pypdf")
        return None
    except Exception as e:
        logger.error(f"Failed to ingest PDF: {e}")
        return None


def ingest_docx(file_path: str) -> Optional[str]:
    """
    Read DOCX file.
    Requires: pip install python-docx
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        logger.info(f"Ingested DOCX: {file_path} ({len(text)} chars from {len(doc.paragraphs)} paragraphs)")
        return text
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return None
    except Exception as e:
        logger.error(f"Failed to ingest DOCX: {e}")
        return None


def batch_ingest(directory: str, pattern: str = "*.*") -> List[Dict[str, Any]]:
    """
    Ingest all documents from a directory.
    
    Args:
        directory: Directory path
        pattern: File pattern (default: all files)
        
    Returns:
        List of {filename, text, file_type} dicts
    """
    directory = Path(directory)
    if not directory.exists():
        logger.error(f"Directory not found: {directory}")
        return []
    
    results = []
    for file_path in directory.glob(pattern):
        if file_path.is_file():
            text = ingest_document(str(file_path))
            if text:
                results.append({
                    'filename': file_path.name,
                    'path': str(file_path),
                    'text': text,
                    'file_type': file_path.suffix.lower(),
                    'size': len(text)
                })
    
    logger.info(f"Batch ingested {len(results)} files from {directory}")
    return results


# ═══════════════════════════════════════════════════════════════════════════════
# KnowledgeIngestor — Phase 2.3 class-based ingestion
# ═══════════════════════════════════════════════════════════════════════════════

import os
import re
import sqlite3
import uuid
from datetime import datetime
from dataclasses import dataclass


@dataclass
class IngestResult:
    chunks_created: int
    vault_file: Optional[str]
    kb_ids: list
    topic: str
    source: str


class KnowledgeIngestor:
    """Pipeline de ingestão de conhecimento — Phase 2.3."""

    def __init__(self) -> None:
        self.db_path = Path(os.getenv("DB_PATH", "database.db"))
        self.vault_path = Path(os.getenv("OBSIDIAN_VAULT_PATH", "./Trade-CLI-Vault"))
        self.chunk_size = 800    # chars por chunk
        self.chunk_overlap = 100  # overlap entre chunks

    def ingest_file(
        self,
        file_path: str,
        topic: str,
        symbol: Optional[str] = None,
        confidence: str = "medium",
    ) -> IngestResult:
        """Ingere um ficheiro PDF, Markdown ou TXT."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Ficheiro não encontrado: {file_path}")

        text = ingest_document(str(path))
        if not text:
            raise ValueError(f"Não foi possível ler: {file_path}")

        return self._process_text(
            text=text,
            source=str(path),
            topic=topic,
            title=path.stem,
            symbol=symbol,
            confidence=confidence,
        )

    def ingest_text(
        self,
        text: str,
        topic: str,
        title: str,
        symbol: Optional[str] = None,
        confidence: str = "medium",
    ) -> IngestResult:
        """Ingere texto directamente."""
        return self._process_text(
            text=text,
            source="manual",
            topic=topic,
            title=title,
            symbol=symbol,
            confidence=confidence,
        )

    def _chunk_text(self, text: str) -> List[str]:
        """Divide texto em chunks semânticos com overlap."""
        paragraphs = re.split(r"\n{2,}", text.strip())
        chunks: List[str] = []
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) < self.chunk_size:
                current += ("\n\n" if current else "") + para
            else:
                if current:
                    chunks.append(current)
                # Overlap: últimas N chars do chunk anterior
                overlap_text = current[-self.chunk_overlap:] if len(current) > self.chunk_overlap else current
                current = overlap_text + "\n\n" + para

        if current:
            chunks.append(current)

        return [c for c in chunks if len(c.strip()) > 50]

    def _process_text(
        self,
        text: str,
        source: str,
        topic: str,
        title: str,
        symbol: Optional[str],
        confidence: str,
    ) -> IngestResult:
        chunks = self._chunk_text(text)
        kb_ids: List[str] = []
        now = datetime.utcnow().isoformat()

        # 1. Guardar no SQLite
        if self.db_path.exists():
            try:
                conn = sqlite3.connect(self.db_path)
                # Check if knowledge_base table exists
                cursor = conn.execute(
                    "SELECT name FROM sqlite_master WHERE type='table' AND name='knowledge_base'"
                )
                if cursor.fetchone():
                    for i, chunk in enumerate(chunks):
                        chunk_id = f"kb_{uuid.uuid4().hex[:12]}"
                        conn.execute(
                            """
                            INSERT OR REPLACE INTO knowledge_base
                            (id, title, content, content_type, topic, symbol,
                             confidence_level, source, tags, review_status, created_at, updated_at)
                            VALUES (?, ?, ?, 'chunk', ?, ?, ?, ?, '[]', 'active', ?, ?)
                            """,
                            (chunk_id, f"{title} (parte {i+1})", chunk, topic,
                             symbol, confidence, source, now, now),
                        )
                        kb_ids.append(chunk_id)
                    conn.commit()
                conn.close()
            except Exception as e:
                logger.warning(f"KB write failed: {e}")

        # 2. Criar nota no Obsidian vault/treino/
        vault_file = None
        treino_dir = self.vault_path / "treino"
        treino_dir.mkdir(exist_ok=True)

        safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-").lower()
        md_path = treino_dir / f"{datetime.utcnow().strftime('%Y%m%d')}-{safe_title}.md"

        frontmatter = f"""---
type: training-note
topic: {topic}
symbol: {symbol or 'all'}
source: {source}
confidence: {confidence}
status: active
created: {now}
tags: [{topic}, {symbol or 'geral'}, treino]
chunks: {len(chunks)}
---

# {title}

> Ingerido em {now[:10]} | {len(chunks)} chunks | Tópico: {topic}

"""
        preview = text[:500].strip() + ("..." if len(text) > 500 else "")
        md_content = frontmatter + preview + "\n"

        try:
            md_path.write_text(md_content, encoding="utf-8")
            vault_file = str(md_path)
        except Exception as e:
            logger.warning(f"Vault write failed: {e}")

        logger.info(f"Knowledge ingested: {title}, {len(chunks)} chunks, topic={topic}")

        return IngestResult(
            chunks_created=len(chunks),
            vault_file=vault_file,
            kb_ids=kb_ids,
            topic=topic,
            source=source,
        )

